"""Day 9 (SQF): Extract PDFs and Word files into page-aware cleaned records.

Output: one JSONL per source document in data/raw/, each line a {page, text}
record. DOCX has no reliable pages, so its page is null. Day 10 (chunk.py)
consumes these.
"""
import json
import re
from collections import Counter
from pathlib import Path

import fitz  # pymupdf
from docx import Document as DocxDocument
from docx.oxml.ns import qn

SOURCE_DIR = Path("data/source")
RAW_DIR = Path("data/raw")

# doc_type and edition are not guessable from the file, so declare them here.
# edition matches YOUR documents: the Spices, Inc. certification report identifies
# this corpus as SQF Fundamentals Edition 1.1 (FSC 19), not SQF Code Edition 9.
_ED = "SQF Fundamentals 1.1"
_ED_CODE10 = "SQF Code 10"   # the full Food Safety Code, a different (newer) standard
MANIFEST = {
    "Certification Report 1.pdf":                     {"doc_type": "report",  "edition": _ED},
    "_Food Safety Management System.pdf":             {"doc_type": "manual",  "edition": _ED},
    # The authoritative published standards themselves (not site SOPs). The site is
    # certified against Fundamentals 1.1; Code Edition 10 is included as reference and
    # kept on its own edition tag so answers do not conflate the two standards.
    "sqf-fundamentals-for-manufacturing-intermediate-09262019-ed-1-1-final.pdf":
                                                      {"doc_type": "standard", "edition": _ED},
    "SQFI-Food-Safety-Code-Edition-10_FSC-19.pdf":    {"doc_type": "standard", "edition": _ED_CODE10},
    "2.1.1 Food Safety Policy.docx":                  {"doc_type": "policy",  "edition": _ED},
    "Food Safety Reporting Structure Statement.docx": {"doc_type": "policy",  "edition": _ED},
    "2.1.2 Management Responsibility.docx":                          {"doc_type": "sop", "edition": _ED},
    "2.1.3 Management Review.docx":                                  {"doc_type": "sop", "edition": _ED},
    "2.1.4 Complaint Management.docx":                               {"doc_type": "sop", "edition": _ED},
    "2.2.1 Food Safety Management System.docx":                      {"doc_type": "sop", "edition": _ED},
    "2.2.2 Document Control.docx":                                   {"doc_type": "sop", "edition": _ED},
    "2.2.3 Records.docx":                                            {"doc_type": "sop", "edition": _ED},
    "2.3.2 Raw and Packaging Materials.docx":                        {"doc_type": "sop", "edition": _ED},
    "2.3.5 Finished Product.docx":                                   {"doc_type": "sop", "edition": _ED},
    "2.4.1 Food Legislation.docx":                                   {"doc_type": "sop", "edition": _ED},
    "2.4.2 Good Manufacturing Practices.docx":                       {"doc_type": "sop", "edition": _ED},
    "2.4.4 Approved Supplier Program.docx":                          {"doc_type": "sop", "edition": _ED},
    "2.4.5 Non-Conforming Material and Product.docx":               {"doc_type": "sop", "edition": _ED},
    "2.4.6 Product Rework.docx":                                     {"doc_type": "sop", "edition": _ED},
    "2.4.7 Product Release.docx":                                    {"doc_type": "sop", "edition": _ED},
    "2.4.8 Environmental Monitoring.docx":                           {"doc_type": "sop", "edition": _ED},
    "2.5.1, 2.5.2 Validation, Effectiveness and Verification.docx":  {"doc_type": "sop", "edition": _ED},
    "2.5.3 Corrective and Preventative Action.docx":                 {"doc_type": "sop", "edition": _ED},
    "2.5.4 Product Sampling, Inspection and Analysis.docx":          {"doc_type": "sop", "edition": _ED},
    "2.5.5 Internal Audits and Inspections.docx":                    {"doc_type": "sop", "edition": _ED},
    "2.6.1, 2.6.2 Product Identification and Trace.docx":            {"doc_type": "sop", "edition": _ED},
    "2.6.3 Withdrawal and Recall.docx":                              {"doc_type": "sop", "edition": _ED},
    "2.7.1 Food Defense Plan.docx":                                  {"doc_type": "sop", "edition": _ED},
    "2.8.1 Allergen Management For Food Fundamentals.docx":          {"doc_type": "sop", "edition": _ED},
    "2.9 Training.docx":                                             {"doc_type": "sop", "edition": _ED},
    "11.1.7 Equipment Utensils.docx":                               {"doc_type": "sop", "edition": _ED},
    "11.2.1 Repairs and Maintenance.docx":                           {"doc_type": "sop", "edition": _ED},
    "11.2.3 Calibration.docx":                                       {"doc_type": "sop", "edition": _ED},
    "11.2.5 Cleaning and Sanitation.docx":                           {"doc_type": "sop", "edition": _ED},
    "11.3.1 Personnel Hygiene and Welfare.docx":                     {"doc_type": "sop", "edition": _ED},
    "11.3.1 Risk Assessment – Personal Items, Jewelry, Electronics and Drink Containers.docx": {"doc_type": "sop", "edition": _ED},
    "11.4.1 Sensory Evaluation Procedure.docx":                      {"doc_type": "sop", "edition": _ED},
    "11.5.1 Water Ice and Air Supply.docx":                          {"doc_type": "sop", "edition": _ED},
    "11.6.1 Receipt, Storage and Handling of Goods.docx":            {"doc_type": "sop", "edition": _ED},
    "11.6.5 Loading Transport and Unloading Practices.docx":         {"doc_type": "sop", "edition": _ED},
    "11.7.3 Foreign Material Control & Detection.docx":              {"doc_type": "sop", "edition": _ED},
    "11.8.1.1 Waste Disposal.docx":                                  {"doc_type": "sop", "edition": _ED},
    "11.8.1.6 Controlled Disposal of Trademarked Materials.docx":    {"doc_type": "sop", "edition": _ED},
    "Process Narratives copy.docx":                                  {"doc_type": "sop", "edition": _ED},
}


def load_pdf(path: Path) -> list[dict]:
    doc = fitz.open(path)
    try:
        return [{"page": i, "text": pg.get_text("text")}
                for i, pg in enumerate(doc, start=1)]
    finally:
        doc.close()


def is_probably_scanned(pages: list[dict], min_chars_per_page: int = 50) -> bool:
    if not pages:
        return False
    return sum(len(p["text"].strip()) for p in pages) / len(pages) < min_chars_per_page


def ensure_text_layer(path: Path) -> Path:
    """OCR a scanned PDF into a searchable copy; return the path to actually read."""
    if not is_probably_scanned(load_pdf(path)):
        return path
    import ocrmypdf
    out = RAW_DIR / f"{path.stem}.ocr.pdf"
    print(f"  {path.name} looks scanned - running OCR")
    ocrmypdf.ocr(path, out, skip_text=True, progress_bar=False)
    return out


def _element_text(element) -> str:
    """Text of a w:p/w:tc in document order, INCLUDING tracked insertions.

    python-docx's Paragraph.text only reads <w:r> runs that are direct children
    of the paragraph, so runs nested in <w:ins> (unaccepted tracked-change
    insertions) are silently dropped - that is how the entire "Recall vs.
    Withdrawal Determination" section of 2.6.3 went missing. We walk every <w:t>
    descendant instead: this picks up <w:ins> text and naturally excludes
    deletions, since deleted text lives in <w:delText>, not <w:t>.
    """
    parts = []
    for node in element.iter():
        if node.tag == qn("w:t"):
            parts.append(node.text or "")
        elif node.tag == qn("w:tab"):
            parts.append("\t")
        elif node.tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
    return "".join(parts)


def load_docx(path: Path) -> list[dict]:
    doc = DocxDocument(path)
    parts = []
    for p in doc.paragraphs:
        text = _element_text(p._p).strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [t for c in row.cells if (t := _element_text(c._tc).strip())]
            if cells:
                parts.append(" | ".join(cells))   # keep tables as pipe-joined rows
    return [{"page": None, "text": "\n".join(parts)}]


_PAGE_NUM = re.compile(r"^\s*(page\s+)?\d+(\s+of\s+\d+)?\s*$", re.IGNORECASE)


def strip_boilerplate(pages: list[dict], threshold: float = 0.6) -> list[dict]:
    """Drop lines that repeat on >= threshold of pages (running headers/footers)."""
    counts: Counter[str] = Counter()
    for p in pages:
        for line in {ln.strip() for ln in p["text"].splitlines() if ln.strip()}:
            counts[line] += 1
    n = max(len(pages), 1)
    boiler = {line for line, c in counts.items() if c / n >= threshold}
    cleaned = []
    for p in pages:
        kept = [ln for ln in p["text"].splitlines()
                if ln.strip() not in boiler and not _PAGE_NUM.match(ln)]
        cleaned.append({**p, "text": "\n".join(kept)})
    return cleaned


def main():
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    for path in sorted(SOURCE_DIR.iterdir()):
        meta = MANIFEST.get(path.name)
        if meta is None:
            print(f"  SKIP {path.name} (not in MANIFEST)")
            continue

        if path.suffix.lower() == ".pdf":
            pages = strip_boilerplate(load_pdf(ensure_text_layer(path)))
        elif path.suffix.lower() in {".docx", ".doc"}:
            pages = load_docx(path)
        else:
            continue

        out = RAW_DIR / f"{path.stem}.jsonl"
        with out.open("w", encoding="utf-8") as f:
            for pg in pages:
                f.write(json.dumps({**pg, "source": path.name, **meta}) + "\n")
        print(f"  {path.name} -> {out.name} ({len(pages)} pages)")


if __name__ == "__main__":
    main()